#!/usr/bin/env python3
"""
Modern, clean GUI for searching and exporting Cursor notes.
Redesigned with sophisticated UI matching modern design standards.
"""

import tkinter as tk
from tkinter import ttk, messagebox, filedialog, font
import json
import re
from pathlib import Path
import os
import time
from datetime import datetime
from note_finder import NoteFinder
from note_parser import NoteParser
from utils.config import Config
import threading

# Import modules for export formats
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    import docx
    from docx.shared import Inches
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

class RoundedFrame(tk.Canvas):
    """A frame with rounded corners"""
    
    def __init__(self, parent, bg="#FFFFFF", width=200, height=100, corner_radius=15, 
                 highlightthickness=0, **kwargs):
        # Get parent background color safely
        try:
            parent_bg = parent.cget('bg')
        except:
            parent_bg = '#FFFFFF'  # Default to white
            
        tk.Canvas.__init__(self, parent, width=width, height=height, 
                          highlightthickness=highlightthickness, bg=parent_bg, **kwargs)
        
        self.bg = bg
        self.corner_radius = corner_radius
        
        # Draw rounded rectangle
        self._draw_rounded_rect()
        
        # Bind configure event to redraw when resized
        self.bind("<Configure>", self._on_resize)
        
    def _draw_rounded_rect(self):
        """Draw the rounded rectangle background"""
        self.delete("rounded_rect")
        
        width = self.winfo_width()
        height = self.winfo_height()
        
        # Create rounded rectangle
        self.create_rounded_rect(0, 0, width, height, 
                               self.corner_radius, fill=self.bg, 
                               outline="", tags="rounded_rect")
    
    def create_rounded_rect(self, x1, y1, x2, y2, radius, **kwargs):
        """Create a rounded rectangle"""
        points = [
            x1 + radius, y1,
            x2 - radius, y1,
            x2, y1,
            x2, y1 + radius,
            x2, y2 - radius,
            x2, y2,
            x2 - radius, y2,
            x1 + radius, y2,
            x1, y2,
            x1, y2 - radius,
            x1, y1 + radius,
            x1, y1
        ]
        
        return self.create_polygon(points, smooth=True, **kwargs)
    
    def _on_resize(self, event):
        """Handle resize event"""
        self._draw_rounded_rect()

class RoundedButton:
    """Custom rounded button using Canvas."""
    
    def __init__(self, parent, text, command, bg="#2D2D2D", fg="#FFFFFF", 
                 hover_bg="#404040", width=120, height=35, corner_radius=8, font_obj=None):
        self.parent = parent
        self.text = text
        self.command = command
        self.bg = bg
        self.fg = fg
        self.hover_bg = hover_bg
        self.width = width
        self.height = height
        self.corner_radius = corner_radius
        self.font_obj = font_obj or ('Segoe UI', 10)
        self.is_hovered = False
        
        # Get parent background color safely
        try:
            parent_bg = parent.cget('bg')
        except:
            parent_bg = '#FFFFFF'  # Default to white
        
        # Create canvas
        self.canvas = tk.Canvas(parent, width=width, height=height, 
                               highlightthickness=0, bd=0, bg=parent_bg)
        
        # Draw button
        self._draw_button()
        
        # Bind events
        self.canvas.bind('<Button-1>', self._on_click)
        self.canvas.bind('<Enter>', self._on_enter)
        self.canvas.bind('<Leave>', self._on_leave)
        
    def _draw_button(self):
        """Draw the rounded rectangle button."""
        self.canvas.delete("all")
        
        # Choose color based on hover state
        current_bg = self.hover_bg if self.is_hovered else self.bg
        
        # Draw rounded rectangle
        self._draw_rounded_rect(0, 0, self.width, self.height, 
                               self.corner_radius, fill=current_bg, outline="")
        
        # Draw text
        self.canvas.create_text(self.width//2, self.height//2, 
                               text=self.text, fill=self.fg, 
                               font=self.font_obj, anchor='center')
    
    def _draw_rounded_rect(self, x1, y1, x2, y2, radius, **kwargs):
        """Draw a rounded rectangle on canvas."""
        points = []
        
        # Top side
        points.extend([x1 + radius, y1])
        points.extend([x2 - radius, y1])
        
        # Top right corner
        points.extend([x2, y1])
        points.extend([x2, y1 + radius])
        
        # Right side
        points.extend([x2, y2 - radius])
        
        # Bottom right corner
        points.extend([x2, y2])
        points.extend([x2 - radius, y2])
        
        # Bottom side
        points.extend([x1 + radius, y2])
        
        # Bottom left corner
        points.extend([x1, y2])
        points.extend([x1, y2 - radius])
        
        # Left side
        points.extend([x1, y1 + radius])
        
        # Top left corner
        points.extend([x1, y1])
        points.extend([x1 + radius, y1])
        
        return self.canvas.create_polygon(points, smooth=True, **kwargs)
    
    def _on_click(self, event):
        """Handle button click."""
        if self.command:
            self.command()
    
    def _on_enter(self, event):
        """Handle mouse enter."""
        self.is_hovered = True
        self._draw_button()
        self.canvas.config(cursor='hand2')
    
    def _on_leave(self, event):
        """Handle mouse leave."""
        self.is_hovered = False
        self._draw_button()
        self.canvas.config(cursor='')
    
    def pack(self, **kwargs):
        """Pack the canvas."""
        self.canvas.pack(**kwargs)
    
    def grid(self, **kwargs):
        """Grid the canvas."""
        self.canvas.grid(**kwargs)

class RoundedDropdown:
    """A dropdown menu with rounded corners."""
    
    def __init__(self, parent, values, variable, command=None, width=100, height=28, 
                corner_radius=10, bg="#FFFFFF", fg="#333333", font_obj=None):
        self.parent = parent
        self.values = values
        self.variable = variable
        self.command = command
        self.width = width
        self.height = height
        self.corner_radius = corner_radius
        self.bg = bg
        self.fg = fg
        self.font_obj = font_obj or ('Segoe UI', 9)
        self.is_dropped = False
        
        # Get parent background color safely
        try:
            parent_bg = parent.cget('bg')
        except:
            parent_bg = '#FFFFFF'  # Default to white
        
        # Create container frame
        self.frame = ttk.Frame(parent)
        
        # Create button to display selected value
        self.container = RoundedFrame(self.frame, bg=self.bg, corner_radius=corner_radius, width=width, height=height)
        self.container.pack(fill=tk.BOTH, expand=True)
        
        # Create display label and dropdown button in container
        self.display = tk.Label(self.container, 
                              textvariable=variable,
                              font=self.font_obj,
                              bg=self.bg,
                              fg=self.fg,
                              padx=10)
        self.display.pack(side=tk.LEFT, fill=tk.Y)
        
        # Dropdown arrow button
        self.arrow = tk.Label(self.container, 
                           text="▼", 
                           font=('Segoe UI', 8),
                           bg=self.bg,
                           fg=self.fg,
                           padx=5)
        self.arrow.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Bind events
        self.container.bind("<Button-1>", self._on_click)
        self.display.bind("<Button-1>", self._on_click)
        self.arrow.bind("<Button-1>", self._on_click)
        
        # Set initial value if provided
        if values and len(values) > 0:
            if not variable.get():
                variable.set(values[0])
    
    def _on_click(self, event):
        """Show dropdown menu when clicked."""
        # Create a toplevel menu
        menu = tk.Menu(self.parent, tearoff=0)
        
        # Add menu items
        for value in self.values:
            menu.add_command(label=value, 
                           command=lambda v=value: self._select_item(v))
        
        # Get container position
        x = self.container.winfo_rootx()
        y = self.container.winfo_rooty() + self.container.winfo_height()
        
        # Show the menu
        menu.post(x, y)
    
    def _select_item(self, value):
        """Handle item selection."""
        self.variable.set(value)
        if self.command:
            self.command(value)
    
    def pack(self, **kwargs):
        """Pack the frame."""
        self.frame.pack(**kwargs)
    
    def grid(self, **kwargs):
        """Grid the frame."""
        self.frame.grid(**kwargs)

class ModernNoteSearchGUI:
    """Modern, clean GUI for searching through Cursor notes."""
    
    # Color scheme - softer, more professional colors
    COLORS = {
        'bg': '#FFFFFF',           # White background
        'card_bg': '#F8F9FA',      # Very light grey for cards (softer)
        'text': '#333333',         # Dark grey text (softer than black)
        'text_secondary': '#757575', # Secondary text
        'text_light': '#9E9E9E',   # Light text for metadata
        'border': '#E0E0E0',       # Light grey borders
        'button_bg': '#333333',    # Dark grey buttons (changed from blue-grey)
        'button_text': '#FFFFFF',  # White button text
        'button_hover': '#505050', # Lighter grey on hover
        'button_secondary': '#F1F3F4', # Light secondary buttons
        'button_secondary_hover': '#E4E8E9',
        'accent': '#5B9BD5',       # Softer blue accent
        'success': '#66BB6A',      # Softer green for success
        'error': '#EF5350',        # Softer red for errors
        'code_bg': '#F8F9FA',      # Light background for code
        'separator': '#EEEEEE',    # Lighter separator lines
        'shadow': '#E0E0E0',       # Shadow color
        'search_border': '#CCCCCC' # Border for search box (more visible)
    }
    
    def __init__(self):
        """Initialize the modern GUI."""
        self.root = tk.Tk()
        self.root.title("Cursor Note Search & Export")
        self.root.geometry("1200x800")
        self.root.minsize(1000, 600)
        
        # Set window background
        self.root.configure(bg=self.COLORS['bg'])
        
        # Data
        self.all_notes = []
        self.filtered_notes = []
        self.selected_note = None
        
        # Get Downloads folder path
        self.downloads_folder = str(Path.home() / "Downloads")
        
        # Configure styles
        self._configure_styles()
        
        # Create UI
        self._create_widgets()
        
        # Center window on screen
        self._center_window()
        
        # Auto-scan for notes on startup
        self._auto_scan_on_startup()
    
    def _configure_styles(self):
        """Configure ttk styles for modern appearance."""
        self.style = ttk.Style()
        
        # Configure fonts
        self.fonts = {
            'default': font.Font(family='Segoe UI', size=10),
            'heading': font.Font(family='Segoe UI', size=16, weight='normal'),
            'subheading': font.Font(family='Segoe UI', size=12, weight='bold'),
            'small': font.Font(family='Segoe UI', size=9),
            'button': font.Font(family='Segoe UI', size=10),
            'code': font.Font(family='Consolas', size=10),
            'metadata': font.Font(family='Segoe UI', size=9, slant='italic')
        }
        
        # Configure button style
        self.style.configure('Modern.TButton',
                           background=self.COLORS['button_bg'],
                           foreground=self.COLORS['button_text'],
                           borderwidth=0,
                           focuscolor='none',
                           padding=(15, 8),
                           font=self.fonts['button'])
        
        self.style.map('Modern.TButton',
                      background=[('active', self.COLORS['button_hover']),
                                ('pressed', self.COLORS['button_hover'])])
        
        # Configure secondary button style
        self.style.configure('Secondary.TButton',
                           background=self.COLORS['bg'],
                           foreground=self.COLORS['text'],
                           borderwidth=1,
                           relief='solid',
                           focuscolor='none',
                           padding=(15, 8),
                           font=self.fonts['button'])
        
        # Configure entry style
        self.style.configure('Modern.TEntry',
                           fieldbackground=self.COLORS['bg'],
                           borderwidth=1,
                           relief='solid',
                           padding=8)
        
        # Configure frame styles
        self.style.configure('Card.TFrame',
                           background=self.COLORS['card_bg'])
        
        self.style.configure('White.TFrame',
                           background=self.COLORS['bg'])
        
        # Configure shadow frame style
        self.style.configure('Shadow.TFrame',
                           background=self.COLORS['bg'])
        
        # Configure scrollbar style
        self.style.configure('Modern.Vertical.TScrollbar',
                           background=self.COLORS['card_bg'],
                           troughcolor=self.COLORS['bg'],
                           borderwidth=0,
                           arrowsize=14)
    
    def _center_window(self):
        """Center the window on screen."""
        self.root.update_idletasks()
        width = self.root.winfo_width()
        height = self.root.winfo_height()
        x = (self.root.winfo_screenwidth() // 2) - (width // 2)
        y = (self.root.winfo_screenheight() // 2) - (height // 2)
        self.root.geometry(f'{width}x{height}+{x}+{y}')
    
    def _create_widgets(self):
        """Create the modern GUI widgets."""
        
        # Main container with padding
        main_container = ttk.Frame(self.root, style='White.TFrame')
        main_container.pack(fill=tk.BOTH, expand=True, padx=40, pady=30)
        
        # Header section
        self._create_header(main_container)
        
        # Content section with two columns
        content_container = ttk.Frame(main_container, style='White.TFrame')
        content_container.pack(fill=tk.BOTH, expand=True, pady=(0, 0))
        
        # Configure grid weights
        content_container.grid_rowconfigure(0, weight=1)
        content_container.grid_columnconfigure(0, weight=3)  # Left panel slightly smaller
        content_container.grid_columnconfigure(1, weight=5)  # Right panel larger
        
        # Left panel - Note list
        self._create_left_panel(content_container)
        
        # Right panel - Note content
        self._create_right_panel(content_container)
    
    def _create_header(self, parent):
        """Create the header section."""
        header_frame = ttk.Frame(parent, style='White.TFrame')
        header_frame.pack(fill=tk.X, pady=(0, 20))
        
        # Title
        title_label = tk.Label(header_frame, 
                              text="Cursor Note Search & Export",
                              font=self.fonts['heading'],
                              bg=self.COLORS['bg'],
                              fg=self.COLORS['text'])
        title_label.pack(side=tk.LEFT)
        
        # Status on the right
        self.status_label = tk.Label(header_frame,
                                   text="Ready",
                                   font=self.fonts['small'],
                                   bg=self.COLORS['bg'],
                                   fg=self.COLORS['text_secondary'])
        self.status_label.pack(side=tk.RIGHT)
        
        # Rescan button using RoundedButton
        rescan_btn = RoundedButton(header_frame,
                                 text="Rescan Databases",
                                 command=self._scan_notes,
                                 bg=self.COLORS['button_bg'],
                                 fg=self.COLORS['button_text'],
                                 hover_bg=self.COLORS['button_hover'],
                                 width=140,
                                 height=35,
                                 corner_radius=18,
                                 font_obj=self.fonts['button'])
        rescan_btn.pack(side=tk.RIGHT, padx=(0, 20))
        
        # Search section - more prominent
        search_frame = ttk.Frame(parent, style='White.TFrame')
        search_frame.pack(fill=tk.X, pady=(15, 25))
        
        # Search label
        search_label = tk.Label(search_frame,
                              text="Search:",
                              font=self.fonts['subheading'],
                              bg=self.COLORS['bg'],
                              fg=self.COLORS['text'])
        search_label.pack(side=tk.LEFT, padx=(0, 15))
        
        # Search entry container with rounded appearance and visible border
        search_container = RoundedFrame(search_frame, bg=self.COLORS['search_border'], corner_radius=20)
        search_container.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 15), pady=2)
        
        # Inner frame for border effect - creating a visible border
        search_inner = RoundedFrame(search_container, bg=self.COLORS['bg'], corner_radius=19)
        search_inner.pack(fill=tk.BOTH, expand=True, padx=2, pady=2)
        
        # Search entry with custom styling
        self.search_var = tk.StringVar()
        self.search_entry = tk.Entry(search_inner,
                                   textvariable=self.search_var,
                                   font=self.fonts['default'],
                                   bg=self.COLORS['bg'],
                                   fg=self.COLORS['text'],
                                   bd=0,
                                   highlightthickness=0)
        self.search_entry.pack(fill=tk.BOTH, expand=True, padx=15, pady=10)
        self.search_entry.bind('<KeyRelease>', self._on_search_change)
        
        # Clear button using RoundedButton
        clear_btn = RoundedButton(search_frame,
                                text="Clear",
                                command=self._clear_search,
                                bg=self.COLORS['button_secondary'],
                                fg=self.COLORS['text'],
                                hover_bg=self.COLORS['button_secondary_hover'],
                                width=80,
                                height=40,
                                corner_radius=20,
                                font_obj=self.fonts['button'])
        clear_btn.pack(side=tk.LEFT)

    def _create_left_panel(self, parent):
        """Create the left panel with note list."""
        # Create frame for shadow effect
        shadow_frame = ttk.Frame(parent, style='White.TFrame')
        shadow_frame.grid(row=0, column=0, sticky='nsew', padx=(5, 20), pady=5)
        
        # Left panel with rounded corners
        left_panel = RoundedFrame(shadow_frame, bg=self.COLORS['card_bg'], corner_radius=15)
        left_panel.pack(fill=tk.BOTH, expand=True, padx=2, pady=2)
        
        # Create shadow effect
        shadow_frame.configure(style='Shadow.TFrame')
        
        # Inner container with padding
        inner_left = ttk.Frame(left_panel, style='Card.TFrame')
        inner_left.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        # Header with sort options
        header_frame = ttk.Frame(inner_left, style='Card.TFrame')
        header_frame.pack(fill=tk.X, pady=(0, 15))
        
        # Found Notes title
        list_header = tk.Label(header_frame,
                             text="Found Notes",
                             font=self.fonts['subheading'],
                             bg=self.COLORS['card_bg'],
                             fg=self.COLORS['text'])
        list_header.pack(side=tk.LEFT)
        
        # Sort options on the right
        sort_container = ttk.Frame(header_frame, style='Card.TFrame')
        sort_container.pack(side=tk.RIGHT)
        
        tk.Label(sort_container,
                text="Sort by:",
                font=self.fonts['small'],
                bg=self.COLORS['card_bg'],
                fg=self.COLORS['text_secondary']).pack(side=tk.LEFT, padx=(0, 5))
        
        # Use RoundedDropdown for sort options
        self.sort_var = tk.StringVar(value="Title")
        sort_menu = RoundedDropdown(sort_container, 
                                  ["Title", "Recently Modified"], 
                                  self.sort_var, 
                                  command=self._apply_sort,
                                  bg=self.COLORS['card_bg'],
                                  fg=self.COLORS['text'],
                                  corner_radius=10)
        sort_menu.pack(side=tk.LEFT)
        
        self.sort_direction = tk.StringVar(value="Ascending")
        direction_menu = RoundedDropdown(sort_container, 
                                       ["Ascending", "Descending"], 
                                       self.sort_direction, 
                                       command=self._apply_sort,
                                       bg=self.COLORS['card_bg'],
                                       fg=self.COLORS['text'],
                                       corner_radius=10)
        direction_menu.pack(side=tk.LEFT, padx=(5, 0))
        
        # Note listbox with custom styling and rounded corners
        list_frame = RoundedFrame(inner_left, bg=self.COLORS['bg'], corner_radius=10, highlightthickness=0)
        list_frame.pack(fill=tk.BOTH, expand=True)
        
        # Create listbox with scrollbar
        scrollbar = tk.Scrollbar(list_frame, orient=tk.VERTICAL, bg=self.COLORS['card_bg'])
        self.note_listbox = tk.Listbox(list_frame,
                                     font=self.fonts['default'],
                                     bg=self.COLORS['bg'],
                                     fg=self.COLORS['text'],
                                     selectbackground=self.COLORS['accent'],
                                     selectforeground=self.COLORS['button_text'],
                                     activestyle='none',
                                     bd=0,
                                     highlightthickness=0,
                                     yscrollcommand=scrollbar.set)
        scrollbar.config(command=self.note_listbox.yview)
        
        self.note_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=8, pady=8)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.note_listbox.bind('<<ListboxSelect>>', self._on_note_select)
        
    def _create_right_panel(self, parent):
        """Create the right panel with note content."""
        # Create frame for shadow effect
        shadow_frame = ttk.Frame(parent, style='White.TFrame')
        shadow_frame.grid(row=0, column=1, sticky='nsew', padx=5, pady=5)
        
        # Right panel with rounded corners
        right_panel = RoundedFrame(shadow_frame, bg=self.COLORS['card_bg'], corner_radius=15)
        right_panel.pack(fill=tk.BOTH, expand=True, padx=2, pady=2)
        
        # Create shadow effect
        shadow_frame.configure(style='Shadow.TFrame')
        
        # Inner container with padding
        inner_right = ttk.Frame(right_panel, style='Card.TFrame')
        inner_right.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        # Header with action buttons
        header_frame = ttk.Frame(inner_right, style='Card.TFrame')
        header_frame.pack(fill=tk.X, pady=(0, 15))
        
        content_label = tk.Label(header_frame,
                               text="Note Content",
                               font=self.fonts['subheading'],
                               bg=self.COLORS['card_bg'],
                               fg=self.COLORS['text'])
        content_label.pack(side=tk.LEFT)
        
        # Action buttons container
        action_frame = ttk.Frame(header_frame, style='Card.TFrame')
        action_frame.pack(side=tk.RIGHT)
        
        # Export all button
        export_all_btn = RoundedButton(action_frame,
                                     text="Export All Filtered",
                                     command=self._export_filtered,
                                     bg=self.COLORS['button_bg'],
                                     fg=self.COLORS['button_text'],
                                     hover_bg=self.COLORS['button_hover'],
                                     width=130,
                                     height=32,
                                     corner_radius=16,
                                     font_obj=self.fonts['button'])
        export_all_btn.pack(side=tk.LEFT, padx=(0, 8))
        
        # Export button
        export_btn = RoundedButton(action_frame,
                                 text="Export",
                                 command=self._export_note,
                                 bg=self.COLORS['button_bg'],
                                 fg=self.COLORS['button_text'],
                                 hover_bg=self.COLORS['button_hover'],
                                 width=80,
                                 height=32,
                                 corner_radius=16,
                                 font_obj=self.fonts['button'])
        export_btn.pack(side=tk.LEFT, padx=(0, 8))
        
        # Copy button
        copy_btn = RoundedButton(action_frame,
                               text="Copy",
                               command=self._copy_note,
                               bg=self.COLORS['button_bg'],
                               fg=self.COLORS['button_text'],
                               hover_bg=self.COLORS['button_hover'],
                               width=80,
                               height=32,
                               corner_radius=16,
                               font_obj=self.fonts['button'])
        copy_btn.pack(side=tk.LEFT)
        
        # Content text area with custom styling
        content_frame = RoundedFrame(inner_right, bg=self.COLORS['bg'], corner_radius=10, highlightthickness=0)
        content_frame.pack(fill=tk.BOTH, expand=True)
        
        # Create text widget with scrollbar
        scrollbar = tk.Scrollbar(content_frame, orient=tk.VERTICAL, bg=self.COLORS['card_bg'])
        self.content_text = tk.Text(content_frame,
                                  font=self.fonts['default'],
                                  bg=self.COLORS['bg'],
                                  fg=self.COLORS['text'],
                                  wrap=tk.WORD,
                                  bd=0,
                                  highlightthickness=0,
                                  padx=15,
                                  pady=15,
                                  state=tk.DISABLED,
                                  yscrollcommand=scrollbar.set)
        scrollbar.config(command=self.content_text.yview)
        
        self.content_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Configure text tags for better formatting
        self._configure_text_tags()
    
    def _configure_text_tags(self):
        """Configure text formatting tags for better display."""
        # Metadata header tag
        self.content_text.tag_configure("metadata_header", 
                                       font=self.fonts['subheading'],
                                       foreground=self.COLORS['text'],
                                       spacing1=10, spacing3=5)
        
        # Metadata tag
        self.content_text.tag_configure("metadata", 
                                       font=self.fonts['metadata'],
                                       foreground=self.COLORS['text_light'],
                                       spacing3=2)
        
        # Separator tag
        self.content_text.tag_configure("separator", 
                                       font=self.fonts['default'],
                                       foreground=self.COLORS['separator'],
                                       spacing1=5, spacing3=10)
        
        # Content header tag
        self.content_text.tag_configure("content_header", 
                                       font=self.fonts['subheading'],
                                       foreground=self.COLORS['text'],
                                       spacing1=15, spacing3=5)
        
        # Code/content tag - removed grey background
        self.content_text.tag_configure("content", 
                                       font=self.fonts['code'],
                                       foreground=self.COLORS['text'],
                                       spacing1=5,
                                       lmargin1=10, lmargin2=10,
                                       rmargin=10)
        
        # Title tag
        self.content_text.tag_configure("title", 
                                       font=self.fonts['heading'],
                                       foreground=self.COLORS['text'],
                                       spacing3=8)
    
    def _process_findings_into_individual_notes(self, findings):
        """
        Process the findings into individual notes.
        
        Args:
            findings: List of note findings from NoteFinder.
            
        Returns:
            List of individual notes with proper titles.
        """
        individual_notes = []
        
        # Get current time for "last modified" timestamp
        current_time = time.time()
        
        for finding in findings:
            # Get the original finding data
            database = finding['database']
            table = finding['table']
            key = finding['key']
            size = finding['size']
            content = finding['content']
            
            # Try to extract timestamp from database file modification time
            db_path = Path(Config.get_default_workspace_path()) / database / "state.vscdb"
            
            # Get file modification time if possible
            try:
                mod_time = os.path.getmtime(db_path)
            except:
                mod_time = current_time  # Use current time as fallback
            
            # Skip empty JSON structures
            if '{"notepads": {}' in content and '"notepadDataVersion": 0' in content:
                continue
                
            # Parse the content into individual notes
            parsed_notes = NoteParser.parse_notes(content)
            
            if parsed_notes:
                # Create individual note entries for each parsed note
                for parsed_note in parsed_notes:
                    title = parsed_note['title']
                    note_content = parsed_note['content']
                    
                    # Only include notes with actual titles
                    if title and title not in ["Untitled Note", "Empty Note", "Unformatted Note"]:
                        individual_notes.append({
                            'database': database,
                            'table': table,
                            'key': key,
                            'original_key': key,
                            'title': title,
                            'size': len(note_content),
                            'content': note_content,
                            'modified': mod_time  # Add modification time
                        })
            
        return individual_notes
    
    def _auto_scan_on_startup(self):
        """Automatically scan for notes when GUI starts."""
        def startup_scan_worker():
            try:
                self.status_label.config(text="Auto-scanning databases...")
                self.root.update()
                
                # Scan for notes
                finder = NoteFinder()
                findings = finder.find_notes()
                
                if findings:
                    # Process findings into individual notes
                    individual_notes = self._process_findings_into_individual_notes(findings)
                    self.all_notes = individual_notes
                    self.filtered_notes = individual_notes  # Set filtered notes same as all notes initially
                    
                    # Apply initial sorting
                    self._apply_sort()
                    
                    # Update the display
                    self._update_note_list()
                    self.status_label.config(text=f"Auto-scan complete: Found {len(individual_notes)} notes")
                else:
                    self.status_label.config(text="No notes found - Database scan complete")
                
            except Exception as e:
                print(f"Auto-scan failed: {e}")
                self.status_label.config(text=f"Scan failed: {str(e)[:50]}...")
        
        # Run in background thread to prevent GUI freezing
        thread = threading.Thread(target=startup_scan_worker, daemon=True)
        thread.start()
    
    def _scan_notes(self):
        """Scan for notes in a background thread."""
        def scan_worker():
            try:
                self.status_label.config(text="Scanning databases...")
                self.root.update()
                
                finder = NoteFinder()
                findings = finder.find_notes()
                
                if findings:
                    # Process findings into individual notes
                    individual_notes = self._process_findings_into_individual_notes(findings)
                    self.all_notes = individual_notes
                    self._update_note_list()
                    
                    self.status_label.config(text=f"Found {len(individual_notes)} notes")
                else:
                    self.status_label.config(text="No notes found")
                
            except Exception as e:
                messagebox.showerror("Error", f"Failed to scan notes: {e}")
                self.status_label.config(text="Scan failed")
        
        # Run in background thread to prevent GUI freezing
        thread = threading.Thread(target=scan_worker, daemon=True)
        thread.start()
    
    def _on_search_change(self, event=None):
        """Handle search text changes."""
        self._filter_notes()
    
    def _clear_search(self):
        """Clear the search."""
        self.search_var.set("")
        self._filter_notes()
    
    def _filter_notes(self):
        """Filter notes based on search term."""
        search_term = self.search_var.get().lower().strip()
        
        # Start with all notes
        self.filtered_notes = self.all_notes
        
        # Apply search filter if there's a search term
        if search_term:
            self.filtered_notes = []
            for note in self.all_notes:
                # Search in title and content
                title = note.get('title', '').lower()
                if search_term in title or search_term in note['content'].lower():
                    self.filtered_notes.append(note)
        
        # Apply current sort after filtering
        self._apply_sort()
        
        # Update the listbox
        self._update_note_list()
    
    def _update_note_list(self):
        """Update the note listbox."""
        self.note_listbox.delete(0, tk.END)
        
        if not hasattr(self, 'filtered_notes') or self.filtered_notes is None:
            self.filtered_notes = self.all_notes
        
        # Ensure we have notes to display
        if not self.filtered_notes:
            self.note_listbox.insert(tk.END, "No notes found")
            return
            
        # Add each note to the listbox
        for note in self.filtered_notes:
            # Create a display string with just the title or a clean preview
            title = note.get('title', '')
            
            if title:
                # If we have a title, just use that
                display = title
            else:
                # Otherwise use the first line of content as title
                content_preview = note['content'].split('\n')[0].strip()
                # Limit length for display
                if len(content_preview) > 60:
                    content_preview = content_preview[:57] + "..."
                display = content_preview
            
            self.note_listbox.insert(tk.END, display)
        
        # Update status
        if hasattr(self, 'search_var') and self.search_var.get().strip():
            self.status_label.config(text=f"Showing {len(self.filtered_notes)} of {len(self.all_notes)} notes")
        else:
            self.status_label.config(text=f"Showing {len(self.all_notes)} notes")
            
        # Force update GUI
        self.root.update_idletasks()
    
    def _on_note_select(self, event):
        """Handle note selection with smooth visual transition."""
        selection = self.note_listbox.curselection()
        if not selection:
            return
        
        index = selection[0]
        if index < len(self.filtered_notes):
            # Save current position for smooth scroll
            old_position = self.content_text.yview()[0]
            
            # Update selected note
            self.selected_note = self.filtered_notes[index]
            
            # Display with smooth transition
            self._display_note_content(old_position)
    
    def _display_note_content(self, old_position=0):
        """Display the selected note content with rich formatting."""
        if not self.selected_note:
            return
        
        note = self.selected_note
        
        # Clear existing content
        self.content_text.config(state=tk.NORMAL)
        self.content_text.delete(1.0, tk.END)
        
        # Title section (if available)
        title = note.get('title', '')
        if title:
            self.content_text.insert(tk.END, f"{title}\n", "title")
        
        # Metadata section
        self.content_text.insert(tk.END, "Note Details\n", "metadata_header")
        
        # Metadata with nice formatting
        original_key = note.get('original_key', note['key'])
        metadata_items = [
            ("Key", original_key),
            ("Database", note['database']),
            ("Table", note['table']),
            ("Size", f"{note['size']} bytes")
        ]
        
        for label, value in metadata_items:
            self.content_text.insert(tk.END, f"{label}: {value}\n", "metadata")
        
        # Separator
        self.content_text.insert(tk.END, "─" * 50 + "\n", "separator")
        
        # Content header
        self.content_text.insert(tk.END, "Content\n", "content_header")
        
        # Note content with code formatting
        content = note['content']
        
        # Split content into lines for better formatting
        lines = content.split('\n')
        formatted_content = ""
        
        for i, line in enumerate(lines):
            # Add line numbers for code-like content if it looks like code
            if self._looks_like_code(content):
                formatted_content += f"{i+1:3d} | {line}\n"
            else:
                formatted_content += f"{line}\n"
        
        self.content_text.insert(tk.END, formatted_content, "content")
        
        # Disable editing
        self.content_text.config(state=tk.DISABLED)
        
        # Smooth scrolling to top with animation
        self._smooth_scroll_to(0, old_position)
    
    def _smooth_scroll_to(self, target_pos, current_pos, steps=10):
        """Smooth scroll animation."""
        if steps <= 0:
            # Final position
            self.content_text.yview_moveto(target_pos)
            return
        
        # Calculate step size
        step = (target_pos - current_pos) / steps
        next_pos = current_pos + step
        
        # Set next position
        self.content_text.yview_moveto(next_pos)
        
        # Schedule next step
        self.root.after(20, lambda: self._smooth_scroll_to(target_pos, next_pos, steps-1))
    
    def _looks_like_code(self, content):
        """Determine if content looks like code based on heuristics."""
        code_indicators = [
            '{', '}', '[', ']', '(', ')',  # Brackets
            '=', ';', ':', '#',             # Common symbols
            'function', 'class', 'def',     # Keywords
            'import', 'export', 'const',    # More keywords
            'var', 'let', 'if', 'else'      # Basic keywords
        ]
        
        # Count indicators
        indicator_count = sum(1 for indicator in code_indicators if indicator in content)
        
        # If we have multiple code indicators and the content is substantial
        return indicator_count >= 3 and len(content) > 50
    
    def _copy_note(self):
        """Copy the selected note to clipboard."""
        if not self.selected_note:
            messagebox.showwarning("Warning", "No note selected")
            return
        
        try:
            self.root.clipboard_clear()
            self.root.clipboard_append(self.selected_note['content'])
            self._show_success_status("Note copied to clipboard!")
        except Exception as e:
            self._show_error_status(f"Failed to copy note: {e}")
    
    def _show_success_status(self, message):
        """Show a success status message."""
        self.status_label.config(text=message, fg=self.COLORS['success'])
        # Reset to normal after 3 seconds
        self.root.after(3000, lambda: self.status_label.config(text="Ready", fg=self.COLORS['text_secondary']))
    
    def _show_error_status(self, message):
        """Show an error status message."""
        self.status_label.config(text=message, fg=self.COLORS['error'])
        # Reset to normal after 5 seconds
        self.root.after(5000, lambda: self.status_label.config(text="Ready", fg=self.COLORS['text_secondary']))
    
    def _show_info_status(self, message):
        """Show an info status message."""
        self.status_label.config(text=message, fg=self.COLORS['accent'])
    
    def _export_note(self):
        """Export the selected note to a file."""
        if not self.selected_note:
            messagebox.showwarning("Warning", "No note selected")
            return
        
        note = self.selected_note
        title = note.get('title', '')
        original_key = note.get('original_key', note['key'])
        
        # Create filename suggestion
        if title:
            filename_suggestion = f"{title.replace(' ', '_')}"
        else:
            filename_suggestion = f"{original_key.replace(' ', '_')}"
        
        # Define file types based on available modules
        filetypes = [(".txt", "*.txt"), (".md", "*.md")]
        if PDF_SUPPORT:
            filetypes.append((".pdf", "*.pdf"))
        if DOCX_SUPPORT:
            filetypes.append((".docx", "*.docx"))
        filetypes.append(("All files", "*.*"))
        
        # Ask for filename
        filename = filedialog.asksaveasfilename(
            title="Export Note",
            initialdir=self.downloads_folder,
            initialfile=filename_suggestion,
            defaultextension=".txt",
            filetypes=filetypes
        )
        
        if not filename:
            return
        
        try:
            # Determine export format based on file extension
            file_ext = os.path.splitext(filename)[1].lower()
            
            if file_ext == ".pdf" and PDF_SUPPORT:
                self._export_as_pdf(filename, note)
            elif file_ext == ".docx" and DOCX_SUPPORT:
                self._export_as_docx(filename, note)
            elif file_ext == ".md":
                self._export_as_markdown(filename, note)
            else:
                # Default to text export
                self._export_as_text(filename, note)
            
            messagebox.showinfo("Success", f"Note exported to: {filename}")
            
            # Open containing folder
            self._open_containing_folder(filename)
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to export note: {e}")
    
    def _export_as_text(self, filename, note):
        """Export note as text file."""
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(f"Cursor Note Export\n")
            f.write(f"Exported: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write("=" * 50 + "\n\n")
            f.write(f"Key: {note.get('original_key', note['key'])}\n")
            if note.get('title'):
                f.write(f"Title: {note['title']}\n")
            f.write(f"Database: {note['database']}\n")
            f.write(f"Size: {note['size']} bytes\n")
            f.write("-" * 30 + "\n\n")
            f.write(note['content'])
    
    def _export_as_pdf(self, filename, note):
        """Export note as PDF file."""
        doc = SimpleDocTemplate(filename, pagesize=letter)
        styles = getSampleStyleSheet()
        
        # Create story elements
        story = []
        
        # Add header
        story.append(Paragraph(f"<b>Cursor Note Export</b>", styles['Title']))
        story.append(Paragraph(f"Exported: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Add metadata
        story.append(Paragraph(f"<b>Key:</b> {note.get('original_key', note['key'])}", styles['Normal']))
        if note.get('title'):
            story.append(Paragraph(f"<b>Title:</b> {note['title']}", styles['Normal']))
        story.append(Paragraph(f"<b>Database:</b> {note['database']}", styles['Normal']))
        story.append(Paragraph(f"<b>Size:</b> {note['size']} bytes", styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Add content (handle line breaks)
        content_lines = note['content'].split('\n')
        for line in content_lines:
            if line.strip():
                story.append(Paragraph(line, styles['Normal']))
            else:
                story.append(Spacer(1, 6))
        
        # Build PDF
        doc.build(story)
    
    def _export_as_docx(self, filename, note):
        """Export note as DOCX file."""
        doc = docx.Document()
        
        # Add title
        doc.add_heading("Cursor Note Export", 0)
        
        # Add metadata
        doc.add_paragraph(f"Exported: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        meta_table = doc.add_table(rows=1, cols=2)
        meta_table.style = 'Table Grid'
        
        header_cells = meta_table.rows[0].cells
        header_cells[0].text = "Property"
        header_cells[1].text = "Value"
        
        # Add key
        row = meta_table.add_row().cells
        row[0].text = "Key"
        row[1].text = note.get('original_key', note['key'])
        
        # Add title if available
        if note.get('title'):
            row = meta_table.add_row().cells
            row[0].text = "Title"
            row[1].text = note['title']
        
        # Add other metadata
        row = meta_table.add_row().cells
        row[0].text = "Database"
        row[1].text = note['database']
        
        row = meta_table.add_row().cells
        row[0].text = "Size"
        row[1].text = f"{note['size']} bytes"
        
        # Add content
        doc.add_heading("Content", level=1)
        
        # Handle line breaks and paragraphs
        content_lines = note['content'].split('\n')
        current_paragraph = None
        
        for line in content_lines:
            if line.strip():
                if current_paragraph is None:
                    current_paragraph = doc.add_paragraph()
                current_paragraph.add_run(line)
            else:
                current_paragraph = None
                doc.add_paragraph()
        
        # Save document
        doc.save(filename)
    
    def _export_as_markdown(self, filename, note):
        """Export note as Markdown file."""
        with open(filename, 'w', encoding='utf-8') as f:
            # Add header with metadata
            f.write(f"# Cursor Note Export\n\n")
            f.write(f"*Exported: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n\n")
            
            # Add metadata table
            f.write("## Metadata\n\n")
            f.write("| Property | Value |\n")
            f.write("| --- | --- |\n")
            f.write(f"| Key | {note.get('original_key', note['key'])} |\n")
            if note.get('title'):
                f.write(f"| Title | {note['title']} |\n")
            f.write(f"| Database | {note['database']} |\n")
            f.write(f"| Size | {note['size']} bytes |\n\n")
            
            # Add content section
            f.write("## Content\n\n")
            
            # Preserve any markdown formatting in the content
            f.write(note['content'])
    
    def _open_containing_folder(self, filename):
        """Open the folder containing the exported file."""
        try:
            folder_path = os.path.dirname(os.path.abspath(filename))
            
            if os.name == 'nt':  # Windows
                os.startfile(folder_path)
            elif os.name == 'posix':  # macOS, Linux
                import subprocess
                if 'darwin' in os.uname().sysname.lower():  # macOS
                    subprocess.call(['open', folder_path])
                else:  # Linux
                    subprocess.call(['xdg-open', folder_path])
        except:
            # Silently fail if we can't open the folder
            pass
    
    def _export_filtered(self):
        """Export all filtered notes to a file."""
        if not self.filtered_notes:
            messagebox.showwarning("Warning", "No notes to export")
            return
        
        # Define file types based on available modules
        filetypes = [(".txt", "*.txt"), (".md", "*.md")]
        if PDF_SUPPORT:
            filetypes.append((".pdf", "*.pdf"))
        if DOCX_SUPPORT:
            filetypes.append((".docx", "*.docx"))
        filetypes.append(("All files", "*.*"))
        
        # Ask for filename
        filename = filedialog.asksaveasfilename(
            title="Export Filtered Notes",
            initialdir=self.downloads_folder,
            initialfile=f"cursor_notes_filtered_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
            defaultextension=".txt",
            filetypes=filetypes
        )
        
        if not filename:
            return
        
        try:
            # Determine export format based on file extension
            file_ext = os.path.splitext(filename)[1].lower()
            
            if file_ext == ".pdf" and PDF_SUPPORT:
                self._export_filtered_as_pdf(filename)
            elif file_ext == ".docx" and DOCX_SUPPORT:
                self._export_filtered_as_docx(filename)
            elif file_ext == ".md":
                self._export_filtered_as_markdown(filename)
            else:
                # Default to text export
                self._export_filtered_as_text(filename)
            
            messagebox.showinfo("Success", f"Exported {len(self.filtered_notes)} notes to: {filename}")
            
            # Open containing folder
            self._open_containing_folder(filename)
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to export notes: {e}")
    
    def _export_filtered_as_text(self, filename):
        """Export filtered notes as text file."""
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(f"Cursor Notes Export (Filtered)\n")
            f.write(f"Exported: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"Search term: '{self.search_var.get()}'\n")
            f.write(f"Notes exported: {len(self.filtered_notes)}\n")
            f.write("=" * 70 + "\n\n")
            
            for i, note in enumerate(self.filtered_notes, 1):
                f.write(f"NOTE #{i}\n")
                f.write(f"Key: {note.get('original_key', note['key'])}\n")
                if note.get('title'):
                    f.write(f"Title: {note['title']}\n")
                f.write(f"Database: {note['database']}\n")
                f.write(f"Size: {note['size']} bytes\n")
                f.write("-" * 40 + "\n")
                f.write(note['content'])
                f.write("\n" + "=" * 70 + "\n\n")
    
    def _export_filtered_as_pdf(self, filename):
        """Export filtered notes as PDF file."""
        doc = SimpleDocTemplate(filename, pagesize=letter)
        styles = getSampleStyleSheet()
        
        # Create story elements
        story = []
        
        # Add header
        story.append(Paragraph(f"<b>Cursor Notes Export (Filtered)</b>", styles['Title']))
        story.append(Paragraph(f"Exported: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
        story.append(Paragraph(f"Search term: '{self.search_var.get()}'", styles['Normal']))
        story.append(Paragraph(f"Notes exported: {len(self.filtered_notes)}", styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Add each note
        for i, note in enumerate(self.filtered_notes, 1):
            # Add separator for notes after the first one
            if i > 1:
                story.append(Spacer(1, 12))
                story.append(Paragraph("=" * 40, styles['Normal']))
                story.append(Spacer(1, 12))
            
            # Add note header
            story.append(Paragraph(f"<b>NOTE #{i}</b>", styles['Heading1']))
            
            # Add metadata
            story.append(Paragraph(f"<b>Key:</b> {note.get('original_key', note['key'])}", styles['Normal']))
            if note.get('title'):
                story.append(Paragraph(f"<b>Title:</b> {note['title']}", styles['Normal']))
            story.append(Paragraph(f"<b>Database:</b> {note['database']}", styles['Normal']))
            story.append(Paragraph(f"<b>Size:</b> {note['size']} bytes", styles['Normal']))
            story.append(Spacer(1, 12))
            
            # Add content (handle line breaks)
            content_lines = note['content'].split('\n')
            for line in content_lines:
                if line.strip():
                    story.append(Paragraph(line, styles['Normal']))
                else:
                    story.append(Spacer(1, 6))
        
        # Build PDF
        doc.build(story)
    
    def _export_filtered_as_docx(self, filename):
        """Export filtered notes as DOCX file."""
        doc = docx.Document()
        
        # Add title
        doc.add_heading("Cursor Notes Export (Filtered)", 0)
        
        # Add metadata
        doc.add_paragraph(f"Exported: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Search term: '{self.search_var.get()}'")
        doc.add_paragraph(f"Notes exported: {len(self.filtered_notes)}")
        doc.add_paragraph()
        
        # Add each note
        for i, note in enumerate(self.filtered_notes, 1):
            # Add separator for notes after the first one
            if i > 1:
                doc.add_paragraph("=" * 40)
                doc.add_paragraph()
            
            # Add note header
            doc.add_heading(f"NOTE #{i}", level=1)
            
            # Add metadata table
            meta_table = doc.add_table(rows=1, cols=2)
            meta_table.style = 'Table Grid'
            
            header_cells = meta_table.rows[0].cells
            header_cells[0].text = "Property"
            header_cells[1].text = "Value"
            
            # Add key
            row = meta_table.add_row().cells
            row[0].text = "Key"
            row[1].text = note.get('original_key', note['key'])
            
            # Add title if available
            if note.get('title'):
                row = meta_table.add_row().cells
                row[0].text = "Title"
                row[1].text = note['title']
            
            # Add other metadata
            row = meta_table.add_row().cells
            row[0].text = "Database"
            row[1].text = note['database']
            
            row = meta_table.add_row().cells
            row[0].text = "Size"
            row[1].text = f"{note['size']} bytes"
            
            # Add content heading
            doc.add_heading("Content", level=2)
            
            # Handle line breaks and paragraphs
            content_lines = note['content'].split('\n')
            current_paragraph = None
            
            for line in content_lines:
                if line.strip():
                    if current_paragraph is None:
                        current_paragraph = doc.add_paragraph()
                    current_paragraph.add_run(line)
                else:
                    current_paragraph = None
                    doc.add_paragraph()
        
        # Save document
        doc.save(filename)
    
    def _export_filtered_as_markdown(self, filename):
        """Export filtered notes as Markdown file."""
        with open(filename, 'w', encoding='utf-8') as f:
            # Add header
            f.write(f"# Cursor Notes Export (Filtered)\n\n")
            f.write(f"*Exported: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n\n")
            f.write(f"**Search term:** '{self.search_var.get()}'\n\n")
            f.write(f"**Notes exported:** {len(self.filtered_notes)}\n\n")
            
            # Add each note
            for i, note in enumerate(self.filtered_notes, 1):
                # Add separator for notes after the first one
                if i > 1:
                    f.write("\n---\n\n")
                
                # Add note header
                f.write(f"## NOTE #{i}\n\n")
                
                # Add metadata table
                f.write("| Property | Value |\n")
                f.write("| --- | --- |\n")
                f.write(f"| Key | {note.get('original_key', note['key'])} |\n")
                if note.get('title'):
                    f.write(f"| Title | {note['title']} |\n")
                f.write(f"| Database | {note['database']} |\n")
                f.write(f"| Size | {note['size']} bytes |\n\n")
                
                # Add content section
                f.write("### Content\n\n")
                f.write(f"{note['content']}\n\n")
    
    def _apply_sort(self, event=None):
        """Sort the note list based on selected criteria."""
        if not hasattr(self, 'filtered_notes') or not self.filtered_notes:
            return
            
        sort_by = self.sort_var.get()
        direction = self.sort_direction.get()
        reverse = (direction == "Descending")
        
        if sort_by == "Title":
            # Sort by title
            self.filtered_notes.sort(key=lambda x: x.get('title', '').lower(), reverse=reverse)
        elif sort_by == "Recently Modified":
            # Sort by modification time
            self.filtered_notes.sort(key=lambda x: x.get('modified', 0), reverse=not reverse)
        
        # Update the displayed list
        self._update_note_list()
    
    def run(self):
        """Start the GUI."""
        self.root.mainloop()

def main():
    """Main function."""
    app = ModernNoteSearchGUI()
    app.run()

if __name__ == '__main__':
    main() 